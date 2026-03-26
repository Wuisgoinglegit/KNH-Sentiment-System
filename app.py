from flask import Flask, render_template, request, redirect, url_for, jsonify, session, make_response
import sqlite3
from datetime import datetime
from werkzeug.security import generate_password_hash, check_password_hash
import random
import os
import io
from dotenv import load_dotenv

import csv
from io import StringIO
from fpdf import FPDF
import textwrap 
import matplotlib
matplotlib.use('Agg') 
import matplotlib.pyplot as plt

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

import smtplib
from email.message import EmailMessage

from sentiment_engine import SentimentEngine
from department_detection import detect_department 

load_dotenv()
app = Flask(__name__)
app.secret_key = 'knh_secure_super_secret_key_2026' 

engine = SentimentEngine()
DB_NAME = "knh_feedback.db"

def init_db():
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS patient_feedback (
            feedback_id INTEGER PRIMARY KEY AUTOINCREMENT,
            raw_text TEXT NOT NULL,
            sentiment_label VARCHAR(10),
            dept_category VARCHAR(100),
            timestamp DATETIME
        )
    ''')
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS admin_users (
            admin_id INTEGER PRIMARY KEY AUTOINCREMENT,
            staff_id VARCHAR(50) UNIQUE NOT NULL,
            password_hash VARCHAR(255) NOT NULL
        )
    ''')
    
    try: cursor.execute("ALTER TABLE admin_users ADD COLUMN email VARCHAR(100)")
    except sqlite3.OperationalError: pass 
        
    try: cursor.execute("ALTER TABLE admin_users ADD COLUMN phone VARCHAR(20)")
    except sqlite3.OperationalError: pass
        
    try: cursor.execute("ALTER TABLE patient_feedback ADD COLUMN urgency_level VARCHAR(10) DEFAULT 'Low'")
    except sqlite3.OperationalError: pass

    conn.commit()
    conn.close()

init_db()

# ROLE-BASED ACCESS CONTROL (RBAC) LOGIC
def get_allowed_departments(staff_id):
    sid = str(staff_id).upper()
    
    # 1. Full Access "Super Users"
    if sid.startswith('ADM') or sid.startswith('QA'):
        return ['All']
        
    # 2. Front Office & Operations (Strictly non-clinical)
    if sid.startswith('REC'):   
        return ['Reception', 'General', 'Outpatient'] 
    if sid.startswith('BIL'):   
        return ['Billing']
        
    # 3. Wide Access Roles (Nurses)
    if sid.startswith('NUR'):
        return ['Ward', 'Emergency', 'Maternity', 'Pediatrics', 'Outpatient', 'ICU']
        
    # 4. Doctors (Specialty + General/On-Call Base)
    doc_base = ['Outpatient', 'Emergency', 'ICU', 'Ward']
    
    if sid.startswith('DOC'):   return ['Outpatient', 'Emergency', 'ICU', 'Ward']
    if sid.startswith('SURG'):  return ['Surgery'] + doc_base
    if sid.startswith('PED'):   return ['Pediatrics'] + doc_base
    if sid.startswith('MAT'):   return ['Maternity'] + doc_base
    if sid.startswith('ONC'):   return ['Oncology'] + doc_base
    if sid.startswith('REN'):   return ['Renal'] + doc_base
    if sid.startswith('DENT'):  return ['Dental'] + doc_base
    
    # 5. Strict Single-Department Roles (Techs & Pharmacists)
    if sid.startswith('PHARM'): return ['Pharmacy']
    if sid.startswith('LAB'):   return ['Laboratory']
    if sid.startswith('RAD'):   return ['Radiology']
    
    return ['Ward']

@app.context_processor
def inject_access():
    if 'staff_id' in session:
        return dict(allowed_depts=get_allowed_departments(session['staff_id']))
    return dict(allowed_depts=[])

@app.after_request
def add_header(response):
    response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
    response.headers["Pragma"] = "no-cache"
    response.headers["Expires"] = "0"
    return response

# PUBLIC ROUTES
@app.route('/')
def home():
    success = request.args.get('success')
    dept_name = request.args.get('dept')
    return render_template('patient feedback.html', success=success, dept_name=dept_name)

@app.route('/submit', methods=['POST'])
def submit_feedback():
    if request.method == 'POST':
        patient_text = request.form.get('feedback_text')
        department_selections = request.form.getlist('department_selection')

        if not department_selections or "General" in department_selections:
            department_result = detect_department(patient_text)
        else:
            department_result = ", ".join(department_selections)

        sentiment_result = engine.predict(patient_text)
        urgency_result = engine.predict_urgency(patient_text, sentiment_result)
        current_time = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO patient_feedback (raw_text, sentiment_label, dept_category, timestamp, urgency_level)
            VALUES (?, ?, ?, ?, ?)
        ''', (patient_text, sentiment_result, department_result, current_time, urgency_result))
        conn.commit()
        conn.close()

        return redirect(url_for('home', success='true', dept=department_result))

# AUTHENTICATION ROUTES
@app.route('/login', methods=['GET', 'POST'])
def login():
    success_msg = None
    if request.args.get('reset') == 'success': success_msg = "Password reset successfully! Please log in."
        
    error_msg = None
    if request.args.get('error') == 'invalid': error_msg = "Invalid Staff ID or Password."

    if request.method == 'POST':
        staff_id = request.form.get('staffId').upper() # Forces input to uppercase
        password = request.form.get('password')
        
        conn = sqlite3.connect(DB_NAME)
        cursor = conn.cursor()
        # Case-insensitive lookup just in case older lowercase users exist
        cursor.execute('SELECT password_hash, staff_id FROM admin_users WHERE UPPER(staff_id) = ?', (staff_id,))
        user_record = cursor.fetchone()
        conn.close()

        if user_record and check_password_hash(user_record[0], password):
            # Save the exact matched ID into session to preserve their original casing if needed, but we check uppercase later
            session['staff_id'] = user_record[1] 
            
            allowed = get_allowed_departments(staff_id)
            if 'All' in allowed:
                return redirect(url_for('dashboard'))
            else:
                return redirect(url_for('department_analysis', dept_name=allowed[0]))
        else:
            return redirect(url_for('login', error='invalid'))
    return render_template('login.html', error_msg=error_msg, success_msg=success_msg)

@app.route('/forgot', methods=['GET', 'POST'])
def forgot_password():
    step = request.args.get('step', 'request')
    error_msg = None

    if request.method == 'POST':
        if step == 'request':
            staff_id = request.form.get('staffId').upper()
            conn = sqlite3.connect(DB_NAME)
            cursor = conn.cursor()
            cursor.execute('SELECT email FROM admin_users WHERE UPPER(staff_id) = ?', (staff_id,))
            user = cursor.fetchone()
            conn.close()

            if user and user[0]:
                code = str(random.randint(100000, 999999))
                session['reset_code'] = code
                session['reset_staff_id'] = staff_id
                
                try:
                    sender_email = os.getenv("SENDER_EMAIL")
                    app_password = os.getenv("EMAIL_APP_PASSWORD")
                    msg = EmailMessage()
                    msg.set_content(f"Hello,\n\nYour KNH Staff Password Reset Code is: {code}\n\nIf you did not request this, please ignore this email.")
                    msg['Subject'] = 'KNH Password Reset Code'
                    msg['From'] = f"KNH System <{sender_email}>"
                    msg['To'] = user[0]
                    server = smtplib.SMTP_SSL('smtp.gmail.com', 465)
                    server.login(sender_email, app_password)
                    server.send_message(msg)
                    server.quit()
                except Exception as e:
                    pass
                return redirect(url_for('forgot_password', step='verify'))
            else:
                error_msg = "Staff ID not found or no email registered."

        elif step == 'verify':
            entered_code = request.form.get('code')
            if entered_code == session.get('reset_code'):
                return redirect(url_for('forgot_password', step='reset'))
            else:
                error_msg = "Invalid verification code."

        elif step == 'reset':
            new_password = request.form.get('new_password')
            hashed_pw = generate_password_hash(new_password)
            conn = sqlite3.connect(DB_NAME)
            cursor = conn.cursor()
            cursor.execute('UPDATE admin_users SET password_hash = ? WHERE UPPER(staff_id) = ?', (hashed_pw, session.get('reset_staff_id')))
            conn.commit()
            conn.close()
            session.pop('reset_code', None)
            session.pop('reset_staff_id', None)
            return redirect(url_for('login', reset='success'))
    return render_template('forgot.html', step=step, error=error_msg)

@app.route('/logout')
def logout():
    session.pop('staff_id', None)
    return redirect(url_for('login'))

# SUPER ADMIN: STAFF MANAGEMENT (FIXED)
@app.route('/manage_users')
def manage_users():
    # FIX: Forces .upper() to ensure admin01 doesn't get kicked out!
    if 'staff_id' not in session or not session['staff_id'].upper().startswith('ADM'):
        return redirect(url_for('dashboard')) 
        
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute("SELECT admin_id, staff_id, email FROM admin_users ORDER BY admin_id DESC")
    users = cursor.fetchall()
    conn.close()
    
    success = request.args.get('success')
    error = request.args.get('error')
    
    return render_template('manage_users.html', users=users, staff_id=session['staff_id'], success=success, error=error)

@app.route('/admin/add_user', methods=['POST'])
def admin_add_user():
    if 'staff_id' not in session or not session['staff_id'].upper().startswith('ADM'):
        return redirect(url_for('dashboard'))
        
    new_staff_id = request.form.get('staffId').upper()
    email = request.form.get('email')
    password = request.form.get('password')
    
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    
    cursor.execute('SELECT * FROM admin_users WHERE UPPER(staff_id) = ?', (new_staff_id,))
    if cursor.fetchone():
        conn.close()
        return redirect(url_for('manage_users', error='exists'))
        
    hashed_pw = generate_password_hash(password)
    cursor.execute('INSERT INTO admin_users (staff_id, email, password_hash) VALUES (?, ?, ?)', (new_staff_id, email, hashed_pw))
    conn.commit()
    conn.close()
    
    return redirect(url_for('manage_users', success='added'))

@app.route('/admin/delete_user/<staff_id_to_delete>', methods=['POST'])
def admin_delete_user(staff_id_to_delete):
    if 'staff_id' not in session or not session['staff_id'].upper().startswith('ADM'):
        return redirect(url_for('dashboard'))
        
    if staff_id_to_delete.upper() == session['staff_id'].upper():
        return redirect(url_for('manage_users', error='self_delete'))
        
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute("DELETE FROM admin_users WHERE UPPER(staff_id) = ?", (staff_id_to_delete.upper(),))
    conn.commit()
    conn.close()
    
    return redirect(url_for('manage_users', success='deleted'))

# DASHBOARD & ANALYSIS ROUTES
@app.route('/dashboard')
def dashboard():
    if 'staff_id' not in session: return redirect(url_for('login'))
    
    allowed = get_allowed_departments(session['staff_id'])
    if 'All' not in allowed:
        return redirect(url_for('department_analysis', dept_name=allowed[0]))

    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute('SELECT timestamp, dept_category, raw_text, sentiment_label, urgency_level FROM patient_feedback ORDER BY timestamp DESC LIMIT 20')
    feedbacks = cursor.fetchall()
    cursor.execute('SELECT COUNT(*) FROM patient_feedback')
    total = cursor.fetchone()[0] or 0
    cursor.execute('SELECT COUNT(*) FROM patient_feedback WHERE sentiment_label = "Positive"')
    pos_count = cursor.fetchone()[0] or 0
    cursor.execute('SELECT COUNT(*) FROM patient_feedback WHERE sentiment_label = "Negative"')
    neg_count = cursor.fetchone()[0] or 0
    neu_count = total - (pos_count + neg_count)
    conn.close()
    pos_percent = round((pos_count / total * 100), 1) if total > 0 else 0
    return render_template('dashboard.html', feedbacks=feedbacks, total=total, pos_percent=pos_percent, pos_count=pos_count, neu_count=neu_count, neg_count=neg_count, staff_id=session['staff_id'])

@app.route('/api/dashboard_data')
def api_dashboard_data():
    if 'staff_id' not in session: return jsonify({'error': 'Unauthorized'}), 401
    
    allowed = get_allowed_departments(session['staff_id'])
    if 'All' not in allowed: return jsonify({'error': 'Forbidden'}), 403

    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute('SELECT COUNT(*) FROM patient_feedback')
    total = cursor.fetchone()[0] or 0
    cursor.execute('SELECT COUNT(*) FROM patient_feedback WHERE sentiment_label = "Positive"')
    pos_count = cursor.fetchone()[0] or 0
    cursor.execute('SELECT COUNT(*) FROM patient_feedback WHERE sentiment_label = "Negative"')
    neg_count = cursor.fetchone()[0] or 0
    neu_count = total - (pos_count + neg_count)
    pos_percent = round((pos_count / total * 100), 1) if total > 0 else 0
    cursor.execute('SELECT timestamp, dept_category, raw_text, sentiment_label, urgency_level FROM patient_feedback ORDER BY timestamp DESC LIMIT 20')
    recent_feedbacks = cursor.fetchall()
    conn.close()
    return jsonify({'total': total, 'pos_percent': pos_percent, 'pos_count': pos_count, 'neg_count': neg_count, 'neu_count': neu_count, 'feedbacks': recent_feedbacks})

@app.route('/analysis/<dept_name>')
def department_analysis(dept_name):
    if 'staff_id' not in session: return redirect(url_for('login'))
    
    allowed = get_allowed_departments(session['staff_id'])
    if 'All' not in allowed and dept_name not in allowed:
        return redirect(url_for('department_analysis', dept_name=allowed[0]))

    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    cursor.execute("SELECT timestamp, dept_category, raw_text, sentiment_label, urgency_level FROM patient_feedback WHERE dept_category LIKE ? ORDER BY timestamp DESC", ('%' + dept_name + '%',))
    dept_feedbacks = cursor.fetchall()
    cursor.execute("SELECT COUNT(*) FROM patient_feedback WHERE dept_category LIKE ?", ('%' + dept_name + '%',))
    total_dept = cursor.fetchone()[0] or 0
    cursor.execute("SELECT COUNT(*) FROM patient_feedback WHERE dept_category LIKE ? AND sentiment_label = 'Positive'", ('%' + dept_name + '%',))
    pos_dept = cursor.fetchone()[0] or 0
    cursor.execute("SELECT COUNT(*) FROM patient_feedback WHERE dept_category LIKE ? AND sentiment_label = 'Negative'", ('%' + dept_name + '%',))
    neg_dept = cursor.fetchone()[0] or 0
    neu_dept = total_dept - (pos_dept + neg_dept)
    conn.close()
    return render_template('analysis.html', dept=dept_name, feedbacks=dept_feedbacks, count=total_dept, issues=neg_dept, pos_count=pos_dept, neu_count=neu_dept, neg_count=neg_dept, staff_id=session['staff_id'])

# EXPORT ROUTES (CSV, PDF, DOCX)
@app.route('/export/csv')
def export_csv():
    if 'staff_id' not in session: return redirect(url_for('login'))
    dept = request.args.get('dept', 'All')
    
    allowed = get_allowed_departments(session['staff_id'])
    if dept == 'All' and 'All' not in allowed: return redirect(url_for('department_analysis', dept_name=allowed[0]))
    if dept != 'All' and 'All' not in allowed and dept not in allowed: return redirect(url_for('department_analysis', dept_name=allowed[0]))

    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    if dept == 'All':
        cursor.execute('SELECT timestamp, dept_category, raw_text, sentiment_label, urgency_level FROM patient_feedback ORDER BY timestamp DESC')
    else:
        cursor.execute('SELECT timestamp, dept_category, raw_text, sentiment_label, urgency_level FROM patient_feedback WHERE dept_category LIKE ? ORDER BY timestamp DESC', ('%' + dept + '%',))
    feedbacks = cursor.fetchall()
    conn.close()
    si = StringIO()
    cw = csv.writer(si)
    cw.writerow(['Date', 'Department', 'Patient Feedback', 'Status', 'Urgency'])
    for row in feedbacks:
        cw.writerow([row[0], row[1], row[2], row[3], row[4]])
    filename = f'KNH_Feedback_Report_{dept.replace(" ", "_")}.csv'
    response = make_response(si.getvalue())
    response.headers.set('Content-Disposition', 'attachment', filename=filename)
    response.headers.set('Content-Type', 'text/csv')
    return response

class KNH_PDF(FPDF):
    def header(self):
        logo_path = os.path.join(app.root_path, 'static', 'kenyatta-national-hospital-seeklogo.png')
        if os.path.exists(logo_path):
            self.image(logo_path, 10, 8, 25)
        self.set_y(12)
        self.set_x(40)
        self.set_font('helvetica', 'B', 14)
        self.cell(0, 6, 'KENYATTA NATIONAL HOSPITAL', ln=True)
        self.set_x(40)
        self.set_font('helvetica', '', 10)
        self.cell(0, 5, 'P.O. Box 20723-00202 Nairobi', ln=True)
        self.set_x(40)
        self.cell(0, 5, 'Tel: 020 2726300, 0709854000 | Email: knhadmin@knh.or.ke', ln=True)
        self.set_draw_color(0, 16, 46) 
        self.set_line_width(1)
        self.line(10, 38, 200, 38)
        self.set_y(45)

    def footer(self):
        self.set_y(-15)
        self.set_font('helvetica', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}/{{nb}}', align='C')

@app.route('/export/pdf')
def export_pdf():
    if 'staff_id' not in session: return redirect(url_for('login'))
    dept = request.args.get('dept', 'All')
    
    allowed = get_allowed_departments(session['staff_id'])
    if dept == 'All' and 'All' not in allowed: return redirect(url_for('department_analysis', dept_name=allowed[0]))
    if dept != 'All' and 'All' not in allowed and dept not in allowed: return redirect(url_for('department_analysis', dept_name=allowed[0]))
    
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    if dept == 'All':
        cursor.execute('SELECT timestamp, dept_category, raw_text, sentiment_label, urgency_level FROM patient_feedback ORDER BY timestamp DESC')
    else:
        cursor.execute('SELECT timestamp, dept_category, raw_text, sentiment_label, urgency_level FROM patient_feedback WHERE dept_category LIKE ? ORDER BY timestamp DESC', ('%' + dept + '%',))
    feedbacks = cursor.fetchall()
    conn.close()

    pos_count = sum(1 for r in feedbacks if r[3] == 'Positive')
    neu_count = sum(1 for r in feedbacks if r[3] == 'Neutral')
    neg_count = sum(1 for r in feedbacks if r[3] == 'Negative')
    
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))
    labels = ['Positive', 'Neutral', 'Negative']
    raw_sizes = [pos_count, neu_count, neg_count]
    colors = ['#4ade80', '#fcd34d', '#f87171']
    
    pie_labels = [l for s, l in zip(raw_sizes, labels) if s > 0]
    pie_sizes = [s for s in raw_sizes if s > 0]
    pie_colors = [c for s, c in zip(raw_sizes, colors) if s > 0]

    if sum(raw_sizes) == 0:
        pie_sizes = [1, 1, 1]
        pie_labels = ['No Data', '', '']
        pie_colors = ['#E2E8F0', '#E2E8F0', '#E2E8F0']
        ax2.set_ylim(0, 1) 
    
    ax1.pie(pie_sizes, labels=pie_labels, colors=pie_colors, autopct='%1.1f%%', startangle=90)
    ax1.axis('equal')
    ax1.set_title('Sentiment Breakdown')
    
    bar_colors = colors if sum(raw_sizes) > 0 else ['#E2E8F0', '#E2E8F0', '#E2E8F0']
    ax2.bar(labels, raw_sizes, color=bar_colors)
    ax2.set_title('Sentiment Volume')

    chart_path = os.path.join(app.root_path, 'static', f'temp_chart_export.png')
    plt.savefig(chart_path, bbox_inches='tight')
    plt.close()

    pdf = KNH_PDF()
    pdf.add_page()
    
    pdf.set_font("helvetica", "B", 16)
    title = f"Patient Feedback Report - {dept} Department" if dept != 'All' else "Patient Feedback Analysis"
    pdf.cell(0, 10, title, align="C")
    pdf.ln(12)

    if os.path.exists(chart_path):
        pdf.image(chart_path, x=10, w=190)
        pdf.ln(5)

    pdf.set_font("helvetica", "B", 9)
    pdf.set_fill_color(226, 232, 240) 
    
    pdf.cell(25, 6, "Date", border=1, fill=True)
    pdf.cell(35, 6, "Department", border=1, fill=True)
    pdf.cell(80, 6, "Feedback", border=1, fill=True)
    pdf.cell(25, 6, "Status", border=1, fill=True)
    pdf.cell(25, 6, "Urgency", border=1, fill=True)
    pdf.ln()

    pdf.set_font("helvetica", "", 8)
    
    for row in feedbacks:
        date_str = row[0][:10]
        dept_str = row[1]
        raw_text = row[2].encode('latin-1', 'ignore').decode('latin-1')
        status = row[3]
        urgency = row[4] if row[4] else 'Low'

        fb_lines = textwrap.wrap(raw_text, width=45)
        if not fb_lines: fb_lines = [""]
        
        dept_lines = textwrap.wrap(dept_str, width=20)
        if not dept_lines: dept_lines = [""]
        
        max_lines = max(len(fb_lines), len(dept_lines))

        for i in range(max_lines):
            t_date = date_str if i == 0 else ""
            t_status = status if i == 0 else ""
            t_urgency = urgency if i == 0 else ""
            t_dept = dept_lines[i] if i < len(dept_lines) else ""
            t_fb = fb_lines[i] if i < len(fb_lines) else ""
            
            b_style = 'LTR' if i == 0 else 'LR'
            if i == max_lines - 1: b_style = 'LTRB' if max_lines == 1 else 'LRB'

            pdf.cell(25, 6, t_date, border=b_style)
            pdf.cell(35, 6, t_dept, border=b_style)
            pdf.cell(80, 6, t_fb, border=b_style)
            pdf.cell(25, 6, t_status, border=b_style)
            pdf.cell(25, 6, t_urgency, border=b_style)
            pdf.ln()

    filename = f'KNH_Feedback_Report_{dept.replace(" ", "_")}.pdf'
    response = make_response(bytes(pdf.output()))
    response.headers.set('Content-Disposition', 'attachment', filename=filename)
    response.headers.set('Content-Type', 'application/pdf')
    return response

@app.route('/export/word')
def export_word():
    if 'staff_id' not in session: return redirect(url_for('login'))
    dept = request.args.get('dept', 'All')
    
    allowed = get_allowed_departments(session['staff_id'])
    if dept == 'All' and 'All' not in allowed: return redirect(url_for('department_analysis', dept_name=allowed[0]))
    if dept != 'All' and 'All' not in allowed and dept not in allowed: return redirect(url_for('department_analysis', dept_name=allowed[0]))
    
    conn = sqlite3.connect(DB_NAME)
    cursor = conn.cursor()
    if dept == 'All':
        cursor.execute('SELECT timestamp, dept_category, raw_text, sentiment_label, urgency_level FROM patient_feedback ORDER BY timestamp DESC')
    else:
        cursor.execute('SELECT timestamp, dept_category, raw_text, sentiment_label, urgency_level FROM patient_feedback WHERE dept_category LIKE ? ORDER BY timestamp DESC', ('%' + dept + '%',))
    feedbacks = cursor.fetchall()
    conn.close()

    pos_count = sum(1 for r in feedbacks if r[3] == 'Positive')
    neu_count = sum(1 for r in feedbacks if r[3] == 'Neutral')
    neg_count = sum(1 for r in feedbacks if r[3] == 'Negative')
    
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(10, 4))
    labels = ['Positive', 'Neutral', 'Negative']
    raw_sizes = [pos_count, neu_count, neg_count]
    colors = ['#4ade80', '#fcd34d', '#f87171']
    
    pie_labels = [l for s, l in zip(raw_sizes, labels) if s > 0]
    pie_sizes = [s for s in raw_sizes if s > 0]
    pie_colors = [c for s, c in zip(raw_sizes, colors) if s > 0]

    if sum(raw_sizes) == 0:
        pie_sizes = [1, 1, 1]
        pie_labels = ['No Data', '', '']
        pie_colors = ['#E2E8F0', '#E2E8F0', '#E2E8F0']
        ax2.set_ylim(0, 1) 
    
    ax1.pie(pie_sizes, labels=pie_labels, colors=pie_colors, autopct='%1.1f%%', startangle=90)
    ax1.axis('equal')
    ax1.set_title('Sentiment Breakdown')
    
    bar_colors = colors if sum(raw_sizes) > 0 else ['#E2E8F0', '#E2E8F0', '#E2E8F0']
    ax2.bar(labels, raw_sizes, color=bar_colors)
    ax2.set_title('Sentiment Volume')

    chart_path = os.path.join(app.root_path, 'static', f'temp_chart_export.png')
    plt.savefig(chart_path, bbox_inches='tight')
    plt.close()

    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(9)
    
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False
    header_table.columns[0].width = Inches(1.2)
    header_table.columns[1].width = Inches(5.0)
    
    cell_logo = header_table.cell(0, 0)
    cell_text = header_table.cell(0, 1)
    
    logo_path = os.path.join(app.root_path, 'static', 'kenyatta-national-hospital-seeklogo.png')
    if os.path.exists(logo_path):
        paragraph = cell_logo.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(logo_path, width=Inches(1.0))
        
    p = cell_text.paragraphs[0]
    run_title = p.add_run("KENYATTA NATIONAL HOSPITAL\n")
    run_title.bold = True
    run_title.font.size = Pt(14)
    run_info = p.add_run("P.O. Box 20723-00202 Nairobi\nTel: 020 2726300, 0709854000 | Email: knhadmin@knh.or.ke")
    run_info.font.size = Pt(10)
    
    line_paragraph = doc.add_paragraph()
    p_format = line_paragraph.paragraph_format
    p_format.space_after = Pt(10)
    pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="18" w:space="1" w:color="00102E"/></w:pBdr>')
    line_paragraph._p.get_or_add_pPr().append(pBdr)
    
    title_text = f"Patient Feedback Report - {dept} Department" if dept != 'All' else "Patient Feedback Analysis"
    title = doc.add_heading(title_text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if os.path.exists(chart_path):
        doc.add_picture(chart_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph() 
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    table.autofit = False
    
    widths = [Inches(0.8), Inches(1.2), Inches(3.0), Inches(0.8), Inches(0.8)]
    for i, col in enumerate(table.columns):
        col.width = widths[i]
        
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Date'
    hdr_cells[1].text = 'Department'
    hdr_cells[2].text = 'Feedback'
    hdr_cells[3].text = 'Status'
    hdr_cells[4].text = 'Urgency'
    
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        shd = parse_xml(r'<w:shd {} w:fill="E2E8F0"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shd)

    for row in feedbacks:
        row_cells = table.add_row().cells
        row_cells[0].text = str(row[0][:10])
        row_cells[1].text = str(row[1])
        row_cells[2].text = str(row[2])
        row_cells[3].text = str(row[3])
        row_cells[4].text = str(row[4] if row[4] else 'Low')

    mem_stream = io.BytesIO()
    doc.save(mem_stream)
    mem_stream.seek(0)
    
    filename = f'KNH_Feedback_Report_{dept.replace(" ", "_")}.docx'
    response = make_response(mem_stream.getvalue())
    response.headers.set('Content-Disposition', 'attachment', filename=filename)
    response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    return response

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)